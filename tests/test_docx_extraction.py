"""Tests for DOCX extraction (P2).

`docx_to_markdown` now tries Docling first, falls back to python-docx.
"""
from __future__ import annotations

import importlib
from pathlib import Path

import pytest

_office = importlib.import_module("my_llm_wiki.detect-office-convert")
_docling = importlib.import_module("my_llm_wiki.extract-with-docling")


def _make_minimal_docx(path: Path) -> None:
    """Build a minimal .docx using python-docx so we don't need a real fixture."""
    from docx import Document
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Some body text.")
    doc.add_heading("Section A", level=2)
    doc.add_paragraph("More content.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "X"
    table.rows[1].cells[1].text = "1"
    doc.save(str(path))


def test_docx_falls_back_to_python_docx_when_docling_missing(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    sample = tmp_path / "fallback.docx"
    _make_minimal_docx(sample)

    # Force the docling-unavailable path
    monkeypatch.setattr(_docling, "is_docling_available", lambda: False)

    text = _office.docx_to_markdown(sample)
    assert "Title" in text
    assert "Section A" in text


def test_docx_uses_docling_when_available(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    if not _docling.is_docling_available():
        pytest.skip("docling not installed")
    sample = tmp_path / "withdocling.docx"
    _make_minimal_docx(sample)

    # Spy: track whether the docling path was taken
    calls: list[Path] = []
    real_extract = _docling.extract_with_docling

    def spy(p):
        calls.append(p)
        return real_extract(p)

    monkeypatch.setattr(_office, "_docling_extract", spy, raising=False)

    text = _office.docx_to_markdown(sample)
    assert "Title" in text
    assert "Section A" in text
    assert "Some body text" in text
    assert len(calls) == 1, "docling extractor should have been called exactly once"


def test_docx_preserves_table_content(tmp_path: Path) -> None:
    sample = tmp_path / "tabletest.docx"
    _make_minimal_docx(sample)
    text = _office.docx_to_markdown(sample)
    # Both backends should at least preserve the cell values somewhere
    assert "Name" in text
    assert "Value" in text
    assert "X" in text
