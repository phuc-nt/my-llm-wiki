# office file conversion and incremental manifest management
from __future__ import annotations

import hashlib
import importlib
import json
from pathlib import Path
from typing import TYPE_CHECKING

from my_llm_wiki.constants import OUTPUT_DIR

# Docling adapter — lazy-loaded module reference. Tests can monkeypatch
# `_docling_extract` to spy on or stub out the integration without touching
# the real Docling import surface.
_docling = importlib.import_module("my_llm_wiki.extract-with-docling")
_docling_extract = _docling.extract_with_docling

# Default manifest path relative to working directory
_DEFAULT_MANIFEST = f"{OUTPUT_DIR}/manifest.json"


# Threshold below which a multi-page PDF is treated as "probably scanned"
# and worth re-running through Docling with OCR enabled.
_SCANNED_TEXT_THRESHOLD = 50
_SCANNED_MIN_PAGES = 2


def _looks_scanned(text: str, page_count: int) -> bool:
    """Heuristic: PDF likely scanned if it has multiple pages but barely any text."""
    if page_count < _SCANNED_MIN_PAGES:
        return False
    return len((text or "").strip()) < _SCANNED_TEXT_THRESHOLD


def _docling_extract_with_ocr(path: Path) -> dict:
    """Run Docling with OCR enabled. Slow — only call after scanned detection."""
    try:
        from docling.document_converter import DocumentConverter, PdfFormatOption
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions

        opts = PdfPipelineOptions()
        opts.do_ocr = True
        converter = DocumentConverter(
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=opts)}
        )
        result = converter.convert(str(path))
        return _docling._normalize(result.document)
    except Exception as exc:
        return {"text": "", "headings": [], "tables": [], "page_count": 0, "error": str(exc)}


def _pypdf_extract(path: Path) -> str:
    """Fallback PDF text extraction via pypdf — no OCR, no layout analysis."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = [page.extract_text() for page in reader.pages]
        return "\n".join(p for p in pages if p)
    except Exception:
        return ""


def extract_pdf_text(path: Path) -> str:
    """Extract text from a PDF file.

    Strategy:
    1. Native Docling (no OCR) — fast, layout-aware
    2. pypdf fallback — when Docling unavailable
    3. Docling with OCR — only when result looks scanned (sparse text on multi-page PDF)
    """
    path = Path(path)
    if not path.exists():
        return ""

    text = ""
    page_count = 0
    if _docling.is_docling_available():
        result = _docling_extract(path)
        if not result.get("error"):
            text = result.get("text") or ""
            page_count = result.get("page_count") or 0

    if not text.strip():
        text = _pypdf_extract(path)

    # Trigger OCR only for clearly-scanned multi-page PDFs
    if _docling.is_docling_available() and _looks_scanned(text, page_count):
        ocr_result = _docling_extract_with_ocr(path)
        if not ocr_result.get("error"):
            ocr_text = ocr_result.get("text") or ""
            if ocr_text.strip():
                return ocr_text

    return text


def docx_to_markdown(path: Path) -> str:
    """Convert a .docx file to markdown text.

    Tries Docling first when available — it preserves layout, tables,
    and multi-column structure better than python-docx. Falls back to
    python-docx so users without `[docling]` extra still get extraction.
    """
    if _docling.is_docling_available():
        result = _docling_extract(Path(path))
        text = result.get("text") or ""
        if not result.get("error") and text.strip():
            return text
    return _legacy_docx_to_markdown(path)


def _legacy_docx_to_markdown(path: Path) -> str:
    """Original python-docx based extraction. Used as fallback when Docling
    is unavailable or fails."""
    try:
        from docx import Document
        doc = Document(str(path))
        lines = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {text}")
            elif style.startswith("List"):
                lines.append(f"- {text}")
            else:
                lines.append(text)
        # Tables
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            lines.extend([header, sep])
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)
    except ImportError:
        return ""
    except Exception:
        return ""


def xlsx_to_markdown(path: Path) -> str:
    """Convert an .xlsx file to markdown text using openpyxl."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sections = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                if all(cell is None for cell in row):
                    continue
                rows.append([str(cell) if cell is not None else "" for cell in row])
            if not rows:
                continue
            sections.append(f"## Sheet: {sheet_name}")
            if len(rows) >= 1:
                header = "| " + " | ".join(rows[0]) + " |"
                sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
                sections.extend([header, sep])
                for row in rows[1:]:
                    sections.append("| " + " | ".join(row) + " |")
        wb.close()
        return "\n".join(sections)
    except ImportError:
        return ""
    except Exception:
        return ""


def convert_office_file(path: Path, out_dir: Path) -> Path | None:
    """Convert a .docx or .xlsx to a markdown sidecar in out_dir.

    Returns the path of the converted .md file, or None if conversion failed
    or the required library is not installed.
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        text = docx_to_markdown(path)
    elif ext == ".xlsx":
        text = xlsx_to_markdown(path)
    else:
        return None

    if not text.strip():
        return None

    out_dir.mkdir(parents=True, exist_ok=True)
    name_hash = hashlib.sha256(str(path.resolve()).encode()).hexdigest()[:8]
    out_path = out_dir / f"{path.stem}_{name_hash}.md"
    out_path.write_text(
        f"<!-- converted from {path.name} -->\n\n{text}",
        encoding="utf-8",
    )
    return out_path


def load_manifest(manifest_path: str = _DEFAULT_MANIFEST) -> dict[str, float]:
    """Load the file modification time manifest from a previous run."""
    try:
        return json.loads(Path(manifest_path).read_text(encoding="utf-8"))
    except Exception:
        return {}


def save_manifest(files: dict[str, list[str]], manifest_path: str = _DEFAULT_MANIFEST) -> None:
    """Save current file mtimes so the next update run can diff against them."""
    manifest: dict[str, float] = {}
    for file_list in files.values():
        for f in file_list:
            try:
                manifest[f] = Path(f).stat().st_mtime
            except OSError:
                pass  # file deleted between detect() and manifest write - skip it
    Path(manifest_path).parent.mkdir(parents=True, exist_ok=True)
    Path(manifest_path).write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def detect_incremental(root: Path, manifest_path: str = _DEFAULT_MANIFEST) -> dict:
    """Like detect(), but returns only new or modified files since the last run.

    Compares current file mtimes against the stored manifest.
    Use for update mode: re-extract only what changed, merge into existing graph.
    """
    import importlib
    _detect_mod = importlib.import_module("my_llm_wiki.detect-files")
    full = _detect_mod.detect(root)
    manifest = load_manifest(manifest_path)

    if not manifest:
        full["incremental"] = True
        full["new_files"] = full["files"]
        full["unchanged_files"] = {k: [] for k in full["files"]}
        full["new_total"] = full["total_files"]
        return full

    new_files: dict[str, list[str]] = {k: [] for k in full["files"]}
    unchanged_files: dict[str, list[str]] = {k: [] for k in full["files"]}

    for ftype, file_list in full["files"].items():
        for f in file_list:
            stored_mtime = manifest.get(f)
            try:
                current_mtime = Path(f).stat().st_mtime
            except Exception:
                current_mtime = 0
            if stored_mtime is None or current_mtime > stored_mtime:
                new_files[ftype].append(f)
            else:
                unchanged_files[ftype].append(f)

    current_files = {f for flist in full["files"].values() for f in flist}
    deleted_files = [f for f in manifest if f not in current_files]

    new_total = sum(len(v) for v in new_files.values())
    full["incremental"] = True
    full["new_files"] = new_files
    full["unchanged_files"] = unchanged_files
    full["new_total"] = new_total
    full["deleted_files"] = deleted_files
    return full
